from fastapi import FastAPI, File, UploadFile, HTTPException, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from pathlib import Path
import os, shutil, subprocess, uuid
from typing import Dict, Any, List
import mammoth
import re

from detector import detect_placeholders
from ai_processor import generate_field_guidance
from docx import Document

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

BASE_DIR = Path(__file__).parent
UPLOAD_DIR = BASE_DIR / "uploads"
OUTPUT_DIR = BASE_DIR / "output"
UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

# ---------- Utilities ----------

def _safe_name(name: str) -> str:
    return name.replace("/", "_").replace("\\", "_")

def docx_to_pdf(docx_path: Path) -> Path:
    """
    Convert DOCX → PDF. Tries docx2pdf; falls back to LibreOffice (soffice).
    Returns the PDF path or raises HTTPException on failure.
    """
    pdf_path = OUTPUT_DIR / (docx_path.stem + ".pdf")
    # 1) Try docx2pdf
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        if pdf_path.exists():
            return pdf_path
    except Exception:
        pass
    # 2) Try LibreOffice
    try:
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(OUTPUT_DIR), str(docx_path)],
            check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE
        )
        if pdf_path.exists():
            return pdf_path
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF conversion failed: {e}")
    raise HTTPException(status_code=500, detail="PDF conversion failed")

def create_marked_html(docx_path: Path, placeholders: List[Dict[str, Any]]) -> str:
    """
    Convert DOCX to HTML and insert unique markers at detected field positions.
    Handles HTML tags that may appear within bracketed fields.
    """
    # Convert to HTML
    with open(docx_path, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        html = result.value
    
    # Sort placeholders by start position (process in document order)
    sorted_placeholders = sorted(placeholders, key=lambda x: x['start'])
    
    # Track which labels we've seen for duplicate handling
    label_counts = {}
    
    # Process each placeholder in order
    for idx, placeholder in enumerate(sorted_placeholders):
        field_type = placeholder['type']
        field_id = placeholder['id']
        marker = f'<span class="field-marker" data-field-id="{field_id}">__MARKER_{field_id}__</span>'
        
        if field_type == 'bracketed':
            # For bracketed fields like [COMPANY NAME] or $[Amount]
            value = placeholder['value']
            
            # Build a pattern that allows HTML tags between any characters
            chars = []
            in_bracket = False
            for char in value:
                if char == '[':
                    chars.append(r'\[(?:<[^>]+>)*')
                    in_bracket = True
                elif char == ']':
                    chars.append(r'(?:<[^>]+>)*\]')
                    in_bracket = False
                elif char == '$':
                    chars.append(r'\$(?:<[^>]+>)*')
                elif in_bracket and char not in ' \t\n':
                    # Allow HTML tags between characters inside brackets
                    chars.append(re.escape(char) + r'(?:<[^>]+>)*')
                else:
                    chars.append(re.escape(char))
            
            pattern = ''.join(chars)
            
            # Try to find and replace
            match = re.search(pattern, html, re.IGNORECASE)
            if match:
                html = html[:match.start()] + marker + html[match.end():]
            else:
                # Fallback: try simple replacement
                if value in html:
                    html = html.replace(value, marker, 1)
                
        elif field_type == 'signature_line':
            # For signature lines like "By: _______"
            label = placeholder['label']
            
            # Track how many times we've seen this label
            if label not in label_counts:
                label_counts[label] = 0
            
            # Find all occurrences of "Label:" followed by whitespace
            escaped_label = re.escape(label)
            # Match label: followed by spaces/tabs/whitespace, allowing HTML tags
            pattern = f"({escaped_label}:)((?:<[^>]+>|\\s)+)"
            
            # Find all matches
            matches = list(re.finditer(pattern, html))
            
            if label_counts[label] < len(matches):
                # Get the nth occurrence (where n = label_counts[label])
                match = matches[label_counts[label]]
                
                # Replace this specific occurrence
                before = html[:match.start()]
                after = html[match.end():]
                
                # Keep the label and colon, replace the whitespace with marker
                replacement = match.group(1) + ' ' + marker
                html = before + replacement + after
                
                label_counts[label] += 1
    
    return html

def fill_docx_placeholders(src_docx: Path, placeholders: List[Dict[str, Any]]) -> Path:
    """
    Create a filled DOCX by replacing:
      - bracketed tokens like [Company], $[Amount]
      - signature lines 'By:' -> 'By: <value>' etc.
    """
    doc = Document(str(src_docx))

    # Build mapping from our list
    bracket_map = {}
    sig_map = {}
    for p in placeholders:
        user_val = (p.get("input") or "").strip()
        if not user_val:
            continue
        if p["type"] == "bracketed":
            bracket_map[p["value"]] = user_val
        else:
            sig_map[f'{p["label"]}:'] = user_val

    def _replace_text_in_run_container(container, replacer):
        for para in container.paragraphs:
            for k, v in list(bracket_map.items()):
                if k in para.text:
                    para.text = para.text.replace(k, v)
            for k, v in list(sig_map.items()):
                if k in para.text:
                    parts = para.text.rsplit(k, 1)
                    para.text = (parts[0] + k + " " + v) if len(parts) == 2 else para.text

        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    _replace_text_in_run_container(cell, replacer)

    _replace_text_in_run_container(doc, None)

    out_path = OUTPUT_DIR / f"{src_docx.stem}.filled.docx"
    doc.save(str(out_path))
    return out_path

# ---------- API ----------

@app.post("/api/upload")
async def upload_document(file: UploadFile = File(...)):
    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files supported")

    safe = _safe_name(file.filename)
    docx_path = UPLOAD_DIR / safe
    with open(docx_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    # Detect placeholders
    results = detect_placeholders(str(docx_path))
    
    # Sort by line number (document order) instead of by type
    results["placeholders"].sort(key=lambda x: x['line'])

    # Generate AI hints
    guidance = generate_field_guidance(safe, results["placeholders"], pdf_path=None)

    attached = 0
    for p in results["placeholders"]:
        field_guidance = guidance.get(p["id"], {})
        p["hint"] = field_guidance.get("micro", "")  # Use micro for the hint
        p["hint_long"] = field_guidance.get("long", "")  # Store long for guided mode
        p["demo_value"] = field_guidance.get("demo", "")  # Store demo for quick-fill
        if p["hint"] or p["hint_long"] or p["demo_value"]:
            attached += 1

    # Generate pre-marked HTML for live preview
    marked_html = None
    try:
        marked_html = create_marked_html(docx_path, results["placeholders"])
    except Exception:
        pass

    return {
        "success": True,
        "filename": safe,
        "placeholders": results["placeholders"],
        "summary": results["summary"],
        "marked_html": marked_html,
        "pdf_url": None,
        "ai_hints_enabled": len(guidance) > 0
    }

@app.post("/api/export")
async def export_filled(payload: Dict[str, Any] = Body(...)):
    filename = payload.get("filename")
    fields = payload.get("fields", [])
    also_pdf = bool(payload.get("also_pdf", True))

    if not filename:
        raise HTTPException(status_code=400, detail="filename is required")

    src = UPLOAD_DIR / _safe_name(filename)
    if not src.exists():
        raise HTTPException(status_code=404, detail="Source DOCX not found")

    # Create filled DOCX
    filled_docx = fill_docx_placeholders(src, fields)

    resp = {"filled_docx_url": f"/api/file/{filled_docx.name}"}

    if also_pdf:
        try:
            filled_pdf = docx_to_pdf(filled_docx)
            resp["filled_pdf_url"] = f"/api/file/{filled_pdf.name}"
        except Exception:
            resp["filled_pdf_url"] = None

    return resp

@app.get("/api/file/{name}")
async def get_file(name: str):
    path = (OUTPUT_DIR / name) if (OUTPUT_DIR / name).exists() else (UPLOAD_DIR / name)
    if not path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(path=path, filename=path.name)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
