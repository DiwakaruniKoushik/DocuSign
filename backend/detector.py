# detector.py
import re
from docx import Document

CONTEXT_CHARS = 500  # left/right window

def detect_placeholders(file_path):
    doc = Document(file_path)
    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)

    para_texts = [p.text or "" for p in all_paragraphs]
    full_text = "\n".join(para_texts)

    bracketed = detect_bracketed(full_text)
    signature = detect_signature_lines(all_paragraphs, para_texts, full_text)

    all_placeholders = bracketed + signature
    all_placeholders.sort(key=lambda x: x['start'])

    # attach unique id, context and a heuristic label_guess
    for idx, p in enumerate(all_placeholders):
        start = max(0, p["start"] - CONTEXT_CHARS)
        end = min(len(full_text), p["end"] + CONTEXT_CHARS)
        snippet = full_text[start:end]
        snippet = re.sub(r"[ \t]+", " ", snippet)
        snippet = re.sub(r"\n{3,}", "\n\n", snippet)
        p["context"] = snippet

        # Unique stable ID per field (no collisions even if the token repeats)
        p["id"] = f"{p['type']}@L{p['line']}@{idx}"

        # Heuristic label guess:
        # 1) If immediately after a '$[____]' there is '("Some Term")' capture that term.
        label_guess = None
        if p["type"] == "bracketed" and p.get("value", "").startswith("$["):
            m = re.search(r'\(\s*the\s+([^)â€â€œ"]+?)\s*\)', snippet, flags=re.IGNORECASE)
            if m:
                label_guess = m.group(1).strip()
        # 2) For signature lines, keep the label itself
        if p["type"] == "signature_line":
            label_guess = p.get("label", None)
        # 3) Fallback: nearest Capitalized Words before the bracket/colon
        if not label_guess:
            m2 = re.search(r'([A-Z][A-Za-z]+(?: [A-Z][A-Za-z]+){0,4})\s*(?:\$?\[\s*[_A-Za-z0-9]*\s*\]|:)\b', snippet)
            if m2:
                label_guess = m2.group(1).strip()

        p["label_guess"] = label_guess or ""

    return {
        "placeholders": all_placeholders,
        "summary": {
            "total": len(all_placeholders),
            "bracketed": len([x for x in all_placeholders if x["type"] == "bracketed"]),
            "signature_lines": len([x for x in all_placeholders if x["type"] == "signature_line"])
        }
    }

def detect_bracketed(full_text):
    results = []
    for match in re.finditer(r"\$?\[\s*[^\]\n]+\s*\]", full_text):
        line = full_text.count("\n", 0, match.start()) + 1
        results.append({
            "type": "bracketed",
            "value": match.group(0),
            "start": match.start(),
            "end": match.end(),
            "line": line
        })
    return results

def detect_signature_lines(paragraphs, para_texts, full_text):
    results = []
    para_offsets = build_offsets(para_texts)

    for para_idx, (para, text) in enumerate(zip(paragraphs, para_texts)):
        if not text:
            continue
        colon_pos = text.rfind(':')
        if colon_pos == -1:
            continue

        after_colon = text[colon_pos + 1:]
        if not after_colon or not all(c in (' ', '\t') for c in after_colon) or '\t' not in after_colon:
            continue

        label = text[:colon_pos].strip()
        if '\n' in label:
            label = label.split('\n')[-1].strip()

        has_underline = any(run.font.underline for run in para.runs)
        padding_start = para_offsets[para_idx] + colon_pos + 1
        padding_end = para_offsets[para_idx] + len(text)
        line = full_text.count("\n", 0, padding_start) + 1

        num_tabs = after_colon.count('\t')
        num_spaces = after_colon.count(' ')

        results.append({
            "type": "signature_line",
            "label": label,
            "start": padding_start,
            "end": padding_end,
            "line": line,
            "metadata": {
                "tabs": num_tabs,
                "spaces": num_spaces,
                "underlined": has_underline
            }
        })
    return results

def build_offsets(para_texts):
    offsets = []
    cur = 0
    for i, text in enumerate(para_texts):
        offsets.append(cur)
        cur += len(text)
        if i != len(para_texts) - 1:
            cur += 1
    return offsets